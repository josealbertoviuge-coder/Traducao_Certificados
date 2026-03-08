from pdf2docx import Converter
from docx import Document
from translator import traduzir


# =========================
# PDF → DOCX
# =========================
def pdf_para_docx(pdf_path, docx_path):

    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()


# =========================
# TRADUZIR DOCX (IN PLACE)
# =========================
def traduzir_docx(docx_entrada, docx_saida):

    doc = Document(docx_entrada)

    # traduz parágrafos
    for p in doc.paragraphs:

        texto = p.text.strip()

        if texto:
            try:
                p.text = traduzir(texto)
            except:
                pass

    # traduz tabelas
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                texto = cell.text.strip()

                if texto:
                    try:
                        cell.text = traduzir(texto)
                    except:
                        pass

    doc.save(docx_saida)

from docx import Document
from translator import traduzir


# =========================
# CRIAR DOCX A PARTIR DO OCR
# =========================
def criar_docx_ocr(texto, saida):

    doc = Document()

    doc.add_heading("ENGLISH TRANSLATION", level=1)

    for linha in texto.split("\n"):

        linha = linha.strip()

        if not linha:
            continue

        try:
            traducao = traduzir(linha)
        except:
            traducao = linha

        doc.add_paragraph(traducao)

    doc.save(saida)
